from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def export_txt(text: str, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8")
    return path


def export_docx(text: str, path: Path, title: str = "扫描文稿") -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.size = Pt(16)
    for block in _paragraphs(text):
        para = doc.add_paragraph(block)
        para.paragraph_format.space_after = Pt(8)
        for run in para.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(11)
    doc.save(str(path))
    return path


def _paragraphs(text: str) -> list[str]:
    chunks: list[str] = []
    buf: list[str] = []
    for line in text.replace("\r\n", "\n").split("\n"):
        if line.strip() == "":
            if buf:
                chunks.append("".join(buf).strip() if _looks_chinese(buf) else " ".join(buf).strip())
                buf = []
            continue
        buf.append(line.strip())
    if buf:
        chunks.append("".join(buf).strip() if _looks_chinese(buf) else " ".join(buf).strip())
    return [c for c in chunks if c]


def _looks_chinese(lines: list[str]) -> bool:
    sample = "".join(lines)
    cjk = sum(1 for ch in sample if "\u4e00" <= ch <= "\u9fff")
    return cjk >= max(4, len(sample) // 8)
